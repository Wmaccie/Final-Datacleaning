import streamlit as st
import pandas as pd
import numpy as np  # np.number is used
import base64
from pathlib import Path
import datetime
# import json # REMOVED - Unused import statement 'import json':7
import urllib.parse
import io  # For creating in-memory files
import traceback  # IMPORT ADDED for Unresolved reference 'traceback'
from typing import Optional, Dict, List, Any, Tuple  # IMPORT ADDED for Optional and other types

# --- Attempt to import python-docx, provide guidance if missing ---
# Initialize names to None to satisfy linters about definition paths
Document, Inches, Pt, WD_ALIGN_PARAGRAPH = None, None, None, None
PYTHON_DOCX_AVAILABLE = False  # Default to False
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # from docx.enum.section import WD_SECTION_START # Was unused, removed
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    # Names remain None, PYTHON_DOCX_AVAILABLE remains False.
    # UI will inform user if they try to use DOCX features.
    print("INFO: python-docx library not found. DOCX export will be disabled.")
except Exception as e_docx_import:  # Catch any other unexpected import error
    print(f"ERROR: Could not import python-docx components: {e_docx_import}")

# --- Page Specific Configuration ---
PAGE_TITLE = "Export Center | ScrubHub Data Suite"
PAGE_ICON = "📤"

# --- Asset Directory ---
try:
    ASSETS_DIR = Path(__file__).resolve().parent.parent / "assets"
except NameError:
    ASSETS_DIR = Path("assets")


# ==================================
# HELPER FUNCTIONS
# ==================================

def get_current_timestamp(fmt: str = "%Y-%m-%d_%H-%M-%S") -> str:
    return datetime.datetime.now().strftime(fmt)


def df_to_csv_bytes(df: pd.DataFrame) -> bytes:
    if not isinstance(df, pd.DataFrame): return b""
    return df.to_csv(index=False).encode('utf-8')


def df_to_sql_inserts(df: pd.DataFrame, table_name: str) -> str:
    if not isinstance(df, pd.DataFrame) or df.empty:
        return "-- No data to convert to SQL INSERT statements.\n"
    table_name_safe = "".join(c if c.isalnum() else "_" for c in str(table_name).strip())
    if not table_name_safe: table_name_safe = "my_table"  # Fallback

    sql_statements = [f"-- SQL statements for table: {table_name_safe}"]
    cols_with_types = []
    sanitized_col_names_sql = []

    for col_name_orig in df.columns:
        col_name_sql = "".join(c if c.isalnum() else "_" for c in str(col_name_orig))
        if not col_name_sql: col_name_sql = f"column_{len(sanitized_col_names_sql)}"  # Handle empty col names
        sanitized_col_names_sql.append(col_name_sql)

        dtype = df[col_name_orig].dtype
        sql_type = "TEXT"  # Default
        if pd.api.types.is_integer_dtype(dtype):
            sql_type = "INTEGER"
        elif pd.api.types.is_float_dtype(dtype):
            sql_type = "REAL"
        elif pd.api.types.is_datetime64_any_dtype(dtype):
            sql_type = "TIMESTAMP"
        elif pd.api.types.is_bool_dtype(dtype):
            sql_type = "BOOLEAN"
        cols_with_types.append(f'    "{col_name_sql}" {sql_type}')

    create_table_sql = f"CREATE TABLE IF NOT EXISTS \"{table_name_safe}\" (\n" + ",\n".join(cols_with_types) + "\n);"
    sql_statements.append(create_table_sql)
    sql_statements.append("\n-- INSERT statements:\n")

    cols_sql_str = ", ".join([f'"{c}"' for c in sanitized_col_names_sql])

    for _, row in df.iterrows():
        values_list = []
        for val in row.values:
            if pd.isna(val):
                values_list.append("NULL")
            elif isinstance(val, (int, float, np.number)):
                values_list.append(str(val))
            else:
                values_list.append(f"'{str(val).replace("'", "''")}'")  # Escape single quotes
        values_sql_str = ", ".join(values_list)
        sql_statements.append(f"INSERT INTO \"{table_name_safe}\" ({cols_sql_str}) VALUES ({values_sql_str});")
    return "\n".join(sql_statements)


def generate_data_quality_report_text_content(
        original_df: Optional[pd.DataFrame],
        cleaned_df: pd.DataFrame,
        quality_score: Optional[float],
        cleaning_steps: Optional[List[Dict[str, Any]]]  # More specific type for cleaning_steps
) -> Dict[str, Any]:  # Return type more specific
    """Generates structured content for the data quality report."""

    # FIX for "This dictionary creation could be rewritten as a dictionary literal:95"
    generated_on_ts = get_current_timestamp('%Y-%m-%d %H:%M:%S')
    quality_score_str = f"{quality_score:.2f}%" if quality_score is not None else "Not Available"
    content: Dict[str, Any] = {
        "header": "DATA QUALITY & CLEANING REPORT",
        "generated_on": generated_on_ts,
        "quality_score": quality_score_str,
        "sections": []
    }

    # Dataset Overview Section
    overview_section: Dict[str, Any] = {"title": "[Dataset Overview]", "details": []}
    if isinstance(original_df, pd.DataFrame) and not original_df.empty:
        overview_section["details"].append(
            f"Original Dataset Shape: {original_df.shape[0]} rows, {original_df.shape[1]} columns")
        overview_section["details"].append(f"Original Missing Values (Total): {original_df.isnull().sum().sum()}")
    else:
        overview_section["details"].append("Original Dataset: Not available or empty.")

    if isinstance(cleaned_df, pd.DataFrame) and not cleaned_df.empty:  # Should always be true if page loads
        overview_section["details"].append(
            f"Cleaned Dataset Shape: {cleaned_df.shape[0]} rows, {cleaned_df.shape[1]} columns")
        overview_section["details"].append(f"Cleaned Missing Values (Total): {cleaned_df.isnull().sum().sum()}")
    else:  # Should not happen if main page guard works
        overview_section["details"].append("Cleaned Dataset: Not available or empty.")

    if isinstance(original_df, pd.DataFrame) and not original_df.empty and \
            isinstance(cleaned_df, pd.DataFrame) and not cleaned_df.empty:
        overview_section["details"].append(
            f"Rows Change: {cleaned_df.shape[0] - original_df.shape[0]:+} (Cleaned vs. Original)")
        overview_section["details"].append(
            f"Columns Change: {cleaned_df.shape[1] - original_df.shape[1]:+} (Cleaned vs. Original)")
    content["sections"].append(overview_section)

    # Cleaning Actions Summary Section
    actions_section: Dict[str, Any] = {"title": "[Cleaning Actions Summary]", "actions": []}
    if cleaning_steps and isinstance(cleaning_steps, list):  # Check if cleaning_steps is a list
        for i, step in enumerate(cleaning_steps):
            if not isinstance(step, dict): continue  # Skip invalid steps
            action_detail: Dict[str, Any] = {"name": step.get('action', 'Unknown Action'), "info": []}
            details = step.get('details', {})
            if isinstance(details, dict):  # Ensure details is a dict
                if 'summary' in details and isinstance(details['summary'], dict):
                    summary = details['summary']
                    action_detail["info"].append(f"Duplicates Removed: {summary.get('duplicates_removed', 'N/A')}")
                    action_detail["info"].append(
                        f"Missing Values Filled: {summary.get('missing_values_filled', 'N/A')}")
                    action_detail["info"].append(f"Type Conversions: {summary.get('type_conversions', 'N/A')}")
                    action_detail["info"].append(f"Rows Affected: {summary.get('rows_affected', 'N/A')}")
                elif 'user_prompt' in details:
                    action_detail["info"].append(f"User Prompt: {details.get('user_prompt', 'N/A')}")
                    action_detail["info"].append(f"AI Response Snippet: {details.get('ai_response', 'N/A')[:100]}...")
                else:
                    for key, val in details.items():
                        action_detail["info"].append(f"{key.replace('_', ' ').title()}: {str(val)[:150]}")
            actions_section["actions"].append(action_detail)
    if not actions_section["actions"]:  # If loop didn't add anything or cleaning_steps was empty
        actions_section["actions"].append(
            {"name": "No specific cleaning actions were logged or available.", "info": []})
    content["sections"].append(actions_section)

    return content


def format_report_content_to_text(report_content: Dict[str, Any]) -> str:
    """Formats the structured report content into a plain text string."""
    # Linter warning "Multi-step list initialization can be replaced with a list literal:151"
    # is a style suggestion. For readability of report generation, .append() is often clearer.
    lines: List[str] = []
    lines.append("======================================")
    lines.append(f"    {report_content.get('header', 'REPORT')}    ")
    lines.append("======================================")
    lines.append(f"Report Generated: {report_content.get('generated_on', '')}")
    lines.append("-" * 38)
    lines.append(f"Overall Data Quality Score: {report_content.get('quality_score', 'N/A')}")
    lines.append("-" * 38)

    for section in report_content.get("sections", []):
        lines.append(f"\n{section.get('title', 'Section')}")
        if "details" in section:
            for detail in section.get("details", []):
                lines.append(f"  {detail}")
        if "actions" in section:
            for i, action in enumerate(section.get("actions", [])):  # Added enumerate for index consistency
                lines.append(f"\n  Action {i + 1}: {action.get('name', '')}")
                for info_line in action.get("info", []):
                    lines.append(f"    - {info_line}")
        lines.append("-" * 38)  # Placed to be after each section's content

    lines.append("\nEnd of Report - Made with ScrubHub")
    lines.append("=" * 38)
    return "\n".join(lines)


def generate_data_quality_report_docx_bytes(
        report_content: Dict[str, Any],
        logo_path: Path
) -> Optional[bytes]:
    if not PYTHON_DOCX_AVAILABLE or not Document:  # Check if Document itself is available
        st.error(
            "The `python-docx` library is not installed or loaded correctly. Please install it (`pip install python-docx`) to enable Word export.")
        return None

    try:
        doc = Document()

        # --- Setup Footer ---
        section = doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        run_footer = footer_paragraph.add_run("Made with ScrubHub")
        if Pt: run_footer.font.size = Pt(9)  # Check if Pt is available
        run_footer.font.name = 'Calibri'
        if WD_ALIGN_PARAGRAPH: footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Add Logo ---
        if logo_path.is_file():
            try:
                # Add logo at the top of the document, aligned right
                # Adding picture can create a new paragraph, so get it.
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if WD_ALIGN_PARAGRAPH else 2  # Fallback alignment if enum not loaded
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(str(logo_path),
                                     width=Inches(1.5) if Inches else 1.5 * 914400)  # Check if Inches is available
            except FileNotFoundError:
                print(f"Warning: Logo file not found at '{logo_path}' for DOCX.")
            except Exception as e_logo_docx:
                print(f"Warning: Could not add logo to DOCX: {e_logo_docx}")

        # --- Report Title ---
        doc.add_heading(report_content.get('header', 'DATA QUALITY & CLEANING REPORT'), level=1)

        # --- Metadata ---
        doc.add_paragraph(f"Report Generated: {report_content.get('generated_on', '')}")
        p_score = doc.add_paragraph()
        p_score.add_run("Overall Data Quality Score: ").bold = True
        p_score.add_run(str(report_content.get('quality_score', 'N/A')))  # Ensure it's a string

        # --- Sections ---
        for section_data in report_content.get("sections", []):
            doc.add_heading(section_data.get('title', 'Section'), level=2)
            if "details" in section_data:
                for detail_line in section_data.get("details", []):
                    doc.add_paragraph(str(detail_line), style='ListBullet')  # Ensure string
            if "actions" in section_data:
                for i, action_data in enumerate(section_data.get("actions", [])):
                    doc.add_paragraph(f"Action {i + 1}: {action_data.get('name', '')}", style='IntenseQuote')
                    for info_line in action_data.get("info", []):
                        p_info = doc.add_paragraph(style='ListBullet')
                        if Inches: p_info.paragraph_format.left_indent = Inches(0.25)  # Smaller indent
                        p_info.add_run(str(info_line))  # Ensure string

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

    except Exception as e_docx_gen:  # Renamed 'e'
        st.error(f"Error generating .docx report: {e_docx_gen}")
        # UNRESOLVED REFERENCE TRACEBACK: ensure traceback is imported
        print(f"DOCX_GENERATION_ERROR: {e_docx_gen}\n{traceback.format_exc()}")
        return None


# ==================================
# PAGE STYLING
# ==================================
def set_export_page_styling(background_image_path: Path):
    # ... (Styling function from previous response - assumed correct and complete)
    # For brevity, I'll use a condensed version of the style from previous answer.
    # Ensure this styling makes text WHITE by default.
    current_bg_image_css = "background-color: #0E0000;"  # Dark red fallback
    if background_image_path.is_file():
        try:
            with open(background_image_path, "rb") as f_bg:
                base64_bg_img = base64.b64encode(f_bg.read()).decode()
            current_bg_image_css = f"""
                background-image: url("data:image/jpeg;base64,{base64_bg_img}");
                background-size: cover; background-position: center;
                background-repeat: no-repeat; background-attachment: fixed;
            """
        except Exception as e_bg_export:
            print(f"ERROR: Loading background image '{background_image_path}' for export page: {e_bg_export}")
    else:
        print(f"WARNING: Background image not found: '{background_image_path}'. Using fallback color.")

    page_styling_css = f"""
        <style>
            .stApp {{ {current_bg_image_css} color: white !important; }}
            h1, h2, h3, h4, h5, h6, p, label, .stMarkdown, .stMetric label, .stMetric div[data-testid="stMetricValue"] {{
                color: white !important; 
            }}
            .stExpander header {{
                color: white !important; background-color: rgba(0, 0, 0, 0.3) !important; 
                border-bottom: 1px solid #B71C1C !important;
            }}
            .stExpander header:hover {{ background-color: rgba(0, 0, 0, 0.5) !important; }}
            .stButton>button {{ /* General buttons */
                color: white !important; background-color: #B71C1C !important; 
                border: 1px solid #D32F2F !important; border-radius: 8px !important;
                padding: 10px 20px !important; font-weight: bold !important;
            }}
            .stButton>button:hover {{ background-color: #D32F2F !important; }}
            .stDownloadButton>button {{ /* Download buttons slightly different */
                background-color: #c0392b !important; border-color: #e74c3c !important;
            }}
            .stDownloadButton>button:hover {{ background-color: #e74c3c !important; }}
            .stTextInput input {{ color: #333333 !important; background-color: #FFFFFF !important; border: 1px solid #B71C1C !important; }}
            .stTextInput label {{ color: white !important; }}
            .report-preview {{
                background-color: rgba(0, 0, 0, 0.6); color: #E0E0E0; padding: 15px;
                border-radius: 8px; border: 1px solid rgba(255, 255, 255, 0.2);
                max-height: 300px; overflow-y: auto; font-family: 'Consolas', 'Courier New', monospace;
                white-space: pre-wrap; font-size: 0.9em;
            }}
            .share-links a {{ color: #FFC107 !important; text-decoration: none; margin-right: 15px; font-weight: bold; }}
            .share-links a:hover {{ text-decoration: underline; color: #FFD54F !important; }}
        </style>
    """
    st.markdown(page_styling_css, unsafe_allow_html=True)


# ==================================
# MAIN EXPORT PAGE LAYOUT
# ==================================
def export_page_layout():
    # ... (Layout from previous response - assumed correct and complete, with unique keys)
    # For brevity, I'm not re-pasting the entire layout function,
    # but it should use the helper functions defined above.
    # Ensure all st.button, st.download_button, st.text_input etc. have unique keys.

    st.markdown(f"<h1 style='text-align: center; color: white;'>{PAGE_ICON} Export Center</h1>", unsafe_allow_html=True)
    st.markdown(
        "<p style='text-align: center; color: #E0E0E0;'>Download your cleaned data, quality reports, and share key insights.</p>",
        unsafe_allow_html=True)
    st.markdown("<hr style='border-top: 1px solid #B71C1C'>", unsafe_allow_html=True)

    if 'df' not in st.session_state or not isinstance(st.session_state.df, pd.DataFrame) or st.session_state.df.empty:
        st.warning("⚠️ No cleaned data available to export. Please process data on the 'Data Cleaning' page first.")
        if st.button("⬅️ Go to Data Cleaning Page", key="export_goto_clean_page_final"):  # Unique key
            try:
                st.switch_page("pages/d_Clean.py")
            except Exception as e_nav_export:
                st.error(f"Navigation failed: {e_nav_export}")
        return

    cleaned_df: pd.DataFrame = st.session_state.df
    original_df_from_ss: Optional[pd.DataFrame] = st.session_state.get('original_df')
    original_df = original_df_from_ss if isinstance(original_df_from_ss,
                                                    pd.DataFrame) else None  # Ensure it's DF or None

    quality_score: Optional[float] = st.session_state.get('data_quality_score')
    cleaning_steps: List[Dict[str, Any]] = st.session_state.get('cleaning_steps', [])
    trash_data_dict: Dict[str, pd.DataFrame] = st.session_state.get('trash_data', {})
    qualogy_logo_path: Path = ASSETS_DIR / "qualogy_logo.png"

    with st.expander("📊 Data Quality & Cleaning Report", expanded=True):
        report_content_structured = generate_data_quality_report_text_content(original_df, cleaned_df, quality_score,
                                                                              cleaning_steps)
        report_text_plain = format_report_content_to_text(report_content_structured)

        st.markdown("<h6>Report Preview (Plain Text)</h6>", unsafe_allow_html=True)
        st.markdown(f"<div class='report-preview'>{report_text_plain}</div>", unsafe_allow_html=True)

        col_rep_dl1, col_rep_dl2 = st.columns(2)
        with col_rep_dl1:
            st.download_button(
                label="📥 Download Report (.txt)", data=report_text_plain.encode('utf-8'),
                file_name=f"DataQualityReport_ScrubHub_{get_current_timestamp()}.txt",
                mime="text/plain", use_container_width=True, key="download_report_txt_final_v2"
            )
        with col_rep_dl2:
            if PYTHON_DOCX_AVAILABLE:
                docx_bytes = generate_data_quality_report_docx_bytes(report_content_structured, qualogy_logo_path)
                if docx_bytes:
                    st.download_button(
                        label="📄 Download Report (.docx)", data=docx_bytes,
                        file_name=f"DataQualityReport_ScrubHub_{get_current_timestamp()}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True, key="download_report_docx_final_v2"
                    )
            else:
                st.info(
                    "For .docx report downloads, please ensure the `python-docx` library is installed (`pip install python-docx`).")

        st.markdown("<h6 style='margin-top:15px;'>Share Report Summary</h6>", unsafe_allow_html=True)
        report_subject = urllib.parse.quote("ScrubHub: Data Quality Report")
        report_body_preview = urllib.parse.quote(
            f"Key insights from ScrubHub Data Quality Report:\nQuality Score: {report_content_structured.get('quality_score', 'N/A')}\nCleaned Shape: {cleaned_df.shape if isinstance(cleaned_df, pd.DataFrame) else 'N/A'}\nGenerated: {report_content_structured.get('generated_on')}\n\n-- Made with ScrubHub --"
        )
        st.markdown(f"""<div class="share-links">
            <a href="mailto:?subject={report_subject}&body={report_body_preview}" target="_blank">📧 Email</a>
            <a href="https://wa.me/?text={report_body_preview}" target="_blank">💬 WhatsApp</a>
            </div>""", unsafe_allow_html=True)

    st.markdown("<hr style='border-top: 1px solid #B71C1C; margin-top:1em; margin-bottom:1em;'>",
                unsafe_allow_html=True)

    with st.expander("💾 Export Cleaned Data", expanded=True):
        st.markdown("#### Preview of Cleaned Data (First 5 Rows)")
        st.dataframe(cleaned_df.head())
        file_ts = get_current_timestamp()

        st.markdown("<h6>CSV Export</h6>", unsafe_allow_html=True)
        csv_bytes = df_to_csv_bytes(cleaned_df)
        st.download_button(
            label="📥 Download Cleaned Data (.csv)", data=csv_bytes,
            file_name=f"cleaned_data_ScrubHub_{file_ts}.csv", mime="text/csv",
            use_container_width=True, key="download_cleaned_csv_final_v2"
        )

        st.markdown("<h6 style='margin-top:15px;'>SQL (INSERT Statements) Export</h6>", unsafe_allow_html=True)
        sql_table_name_cleaned = st.text_input("Enter Table Name for Cleaned Data SQL:", value="scrubhub_cleaned_data",
                                               key="sql_table_name_cleaned_input_final_v2")
        if st.button("⚙️ Generate SQL for Cleaned Data", key="generate_sql_cleaned_btn_final_v2",
                     use_container_width=True):
            if sql_table_name_cleaned:
                sql_script_cleaned = df_to_sql_inserts(cleaned_df, sql_table_name_cleaned)
                st.session_state.sql_script_cleaned_to_download = sql_script_cleaned
            else:
                st.warning("Please enter a table name for SQL export.")

        if 'sql_script_cleaned_to_download' in st.session_state and st.session_state.sql_script_cleaned_to_download:
            st.download_button(
                label="📥 Download Cleaned Data (.sql)",
                data=st.session_state.sql_script_cleaned_to_download.encode('utf-8'),
                file_name=f"cleaned_data_{sql_table_name_cleaned}_{file_ts}.sql", mime="application/sql",
                use_container_width=True, key="download_cleaned_sql_final_btn_v2"
            )
            st.text_area("SQL Script Preview (Cleaned Data)", st.session_state.sql_script_cleaned_to_download,
                         height=150, key="sql_preview_cleaned_final_v2",
                         help="Includes basic CREATE TABLE and INSERT statements.")

    st.markdown("<hr style='border-top: 1px solid #B71C1C; margin-top:1em; margin-bottom:1em;'>",
                unsafe_allow_html=True)

    with st.expander("🗑️ Export Removed/Altered Data (if captured)", expanded=False):
        if not trash_data_dict:
            st.info("No 'trash' data segments were explicitly captured during cleaning or passed to this page.")
        else:
            st.markdown("The following data segments were captured as 'trash' or significantly altered items:")
            for i, (trash_name, trash_df_item) in enumerate(trash_data_dict.items()):  # Added enumerate for unique keys
                if isinstance(trash_df_item, pd.DataFrame) and not trash_df_item.empty:
                    st.markdown(f"<h6>{trash_name.replace('_', ' ').title()} (Rows: {len(trash_df_item)})</h6>",
                                unsafe_allow_html=True)
                    st.dataframe(trash_df_item.head(3))
                    col_trash_dl1, col_trash_dl2 = st.columns(2)
                    with col_trash_dl1:
                        st.download_button(
                            label=f"📥 {trash_name} (.csv)", data=df_to_csv_bytes(trash_df_item),
                            file_name=f"trash_{trash_name}_{file_ts}.csv", mime="text/csv",
                            key=f"download_trash_{trash_name}_csv_final_v2_{i}", use_container_width=True
                        )

                    sql_table_name_trash_item = f"scrubhub_trash_{trash_name}"
                    sql_script_trash_item = df_to_sql_inserts(trash_df_item, sql_table_name_trash_item)
                    with col_trash_dl2:
                        st.download_button(
                            label=f"📥 {trash_name} (.sql)", data=sql_script_trash_item.encode('utf-8'),
                            file_name=f"trash_{sql_table_name_trash_item}_{file_ts}.sql", mime="application/sql",
                            key=f"download_trash_{trash_name}_sql_final_btn_v2_{i}", use_container_width=True
                        )
                    if i < len(trash_data_dict) - 1:  # Add separator if not the last item
                        st.markdown("---")

    st.markdown("<hr style='border:1px solid #D32F2F; margin-top: 30px; margin-bottom: 15px;'>", unsafe_allow_html=True)
    st.markdown("<p style='text-align:center; color: #E0E0E0;'><i>Thank you for using the ScrubHub Data Suite!</i></p>",
                unsafe_allow_html=True)


# ==================================
# SCRIPT EXECUTION
# ==================================
if __name__ == "__main__":
    st.set_page_config(page_title=PAGE_TITLE, page_icon=PAGE_ICON, layout="wide")
    set_export_page_styling(ASSETS_DIR / "background_blue.jpg")

    if 'df' not in st.session_state:
        st.session_state.df = pd.DataFrame({'A_clean': [1, 2], 'B_clean': ['x', 'y']})
    if 'original_df' not in st.session_state:
        st.session_state.original_df = pd.DataFrame(
            {'A_orig': [1, 2, 3], 'B_orig': ['x', 'y', 'z'], 'C_orig': [True, False, True]})
    if 'data_quality_score' not in st.session_state:
        st.session_state.data_quality_score = 92.5
    if 'cleaning_steps' not in st.session_state:
        st.session_state.cleaning_steps = [{'action': 'Removed 1 duplicate row', 'details': {'rows_affected': 1}}]
    if 'trash_data' not in st.session_state:
        st.session_state.trash_data = {
            'removed_duplicates_example': pd.DataFrame({'A_orig': [3], 'B_orig': ['z'], 'C_orig': [True]}),
            'dropped_low_quality_example': pd.DataFrame({'X': [99], 'Y': ['test']})
        }
    if not (ASSETS_DIR / "qualogy_logo.png").is_file():
        st.warning(
            f"Test Warning: Qualogy logo not found at expected path: {ASSETS_DIR / 'qualogy_logo.png'}. DOCX logo will be missing if run directly without it.")

    export_page_layout()